import os
import pandas as pd
import uuid
from datetime import datetime
from typing import List, Dict, Any, Optional
import re
import numpy as np
import json
import requests
import time
import warnings

# Additional imports for file processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    warnings.warn("PyPDF2 not available. PDF processing disabled.")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    warnings.warn("python-docx not available. DOCX processing disabled.")

from config import DATA_DIR, TRUSTED_FLAG_DEFAULT, TEXT_CHUNK_SIZE, TEXT_CHUNK_OVERLAP, OLLAMA_BASE_URL, SLM_PARSE_MODEL
from utils import logger, log_process_completion

class Document:
    def __init__(self, doc_id: str, file_id: str, file_name: str, doc_type: str, content: str, metadata: Dict[str, Any]):
        self.doc_id = doc_id
        self.file_id = file_id
        self.file_name = file_name
        self.doc_type = doc_type
        self.content = content
        self.metadata = metadata

    def to_dict(self):
        return {
            "doc_id": self.doc_id,
            "file_id": self.file_id,
            "file_name": self.file_name,
            "doc_type": self.doc_type,
            "content": self.content,
            "metadata": self.metadata,
        }

def _get_slm_summary(row_data: Dict[str, Any], column_types: Dict[str, str]) -> str:
    """
    Generates a concise, natural language summary of a single row using an SLM.
    Returns clean, direct summaries without prefixes or verbose explanations.
    """
    # Filter out numeric fields that are NaN for better summary generation
    filtered_row_data = {k: v for k, v in row_data.items() if pd.notna(v)}

    # Create a more direct prompt that focuses on clean output
    prompt = f"""Generate a concise summary of this data row. Return ONLY the summary text, no prefixes like "Summary:" or "Based on the given row".

Examples:
Input: {{'feeder': 'I/C Panel Numerical Relay', 'swb_no': 'I/C-1', '30-06-2024': 246740, '01-07-2024': 246885, 'difference': 145}}
Output: I/C Panel Numerical Relay feeder SWB I/C-1 shows 246740 KWH on 30-06-2024, 246885 KWH on 01-07-2024 with 145 KWH difference.

Input: {{'product': 'Laptop', 'price': 1200, 'category': 'Electronics'}}
Output: Electronics Laptop priced at $1200.

Data: {filtered_row_data}
Summary:"""

    logger.info(f"Requesting SLM summary for row (first 100 chars): {str(filtered_row_data)[:100]}...")
    start_time = time.time()
    try:
        response = requests.post(
            f"{OLLAMA_BASE_URL}/api/generate",
            json={
                "model": SLM_PARSE_MODEL,
                "prompt": prompt,
                "stream": False,
                "options": {"num_gpu": 1, "temperature": 0.1, "top_p": 0.9}
            },
            timeout=30  # Reduced timeout for faster processing
        )
        end_time = time.time()
        response.raise_for_status()
        response_data = response.json()
        logger.info(f"Ollama SLM response status: {response.status_code}, time: {end_time - start_time:.2f}s")
        
        if "response" in response_data and response_data["response"]:
            summary = response_data["response"].strip()
            
            # Clean up common prefixes and verbose additions
            prefixes_to_remove = [
                "summary:",
                "based on the given row of data, here is a concise and natural language summary:",
                "based on the given row,",
                "here is a concise summary:",
                "the summary is:",
                "summary of the data:",
                "data summary:",
                "row summary:",
                "concise summary:",
                "natural language summary:"
            ]
            
            summary_lower = summary.lower()
            for prefix in prefixes_to_remove:
                if summary_lower.startswith(prefix):
                    summary = summary[len(prefix):].strip()
                    break
            
            # Remove any remaining verbose additions
            if "based on" in summary_lower or "given row" in summary_lower:
                # Extract the actual summary part after common verbose phrases
                lines = summary.split('\n')
                for line in lines:
                    line = line.strip()
                    if line and not any(phrase in line.lower() for phrase in ["based on", "given row", "here is", "summary:", "the data"]):
                        summary = line
                        break
            
            logger.info(f"Cleaned SLM summary: {summary[:100]}...")
            return summary
        else:
            logger.warning(f"Ollama SLM returned no 'response' key or it was empty for row: {str(filtered_row_data)[:100]}...")
            logger.warning(f"Full Ollama SLM response: {json.dumps(response_data)}")
            return str(filtered_row_data) # Fallback to string representation

    except requests.exceptions.RequestException as e:
        end_time = time.time()
        logger.error(f"Error getting SLM summary from Ollama (took {end_time - start_time:.2f}s): {e}", exc_info=True)
        if hasattr(e, 'response') and e.response is not None:
            logger.error(f"Ollama SLM error response status: {e.response.status_code}")
            logger.error(f"Ollama SLM error response body: {e.response.text}")
        return str(filtered_row_data) # Fallback to string representation

def _smart_excel_processing(file_path: str, sheet_name: str, row_limit: int = None) -> pd.DataFrame:
    """Smart Excel processing with intelligent header detection and structure analysis"""
    try:
        # Step 1: Analyze the raw structure
        df_raw = pd.read_excel(file_path, sheet_name=sheet_name, header=None, nrows=10)
        structure_analysis = _analyze_excel_structure(df_raw)
        
        logger.info(f"Excel structure analysis for {sheet_name}: {structure_analysis}")
        
        # Step 2: Choose the best processing method
        if structure_analysis['complexity_score'] > 5:
            # Complex structure - use advanced processing
            df = _process_complex_excel_structure(file_path, sheet_name, structure_analysis)
        elif structure_analysis['has_multi_level_headers']:
            # Multi-level headers
            df = pd.read_excel(file_path, sheet_name=sheet_name, header=structure_analysis['header_rows'])
            df.columns = _flatten_multiindex_columns(df.columns)
        else:
            # Simple structure
            df = pd.read_excel(file_path, sheet_name=sheet_name, header=structure_analysis['header_rows'][0])
            df.columns = _normalize_headers(df.columns)
        
        # Step 3: Clean and validate
        df = _clean_dataframe(df)
        
        # Step 4: Apply row limit if specified
        if row_limit and len(df) > row_limit:
            df = df.head(row_limit)
        
        logger.info(f"Successfully processed {sheet_name}: {len(df)} rows, {len(df.columns)} columns")
        logger.info(f"Sample columns: {list(df.columns)[:5]}")
        
        return df
        
    except Exception as e:
        logger.error(f"Smart Excel processing failed for {sheet_name}: {e}")
        # Fallback to simple processing
        try:
            df = pd.read_excel(file_path, sheet_name=sheet_name, header=0)
            df.columns = _normalize_headers(df.columns)
            return df.head(row_limit) if row_limit else df
        except Exception as e2:
            logger.error(f"Fallback processing also failed: {e2}")
            return None

def _analyze_excel_structure(df_raw: pd.DataFrame) -> dict:
    """Analyze Excel structure to determine complexity and processing strategy"""
    analysis = {
        'header_rows': [],
        'data_start_row': 0,
        'has_multi_level_headers': False,
        'has_merged_cells': False,
        'complexity_score': 0,
        'recommended_method': 'simple'
    }
    
    # Analyze first 8 rows
    for i in range(min(8, len(df_raw))):
        row = df_raw.iloc[i]
        non_null_count = row.notna().sum()
        
        if non_null_count == 0:
            continue
        
        # Count text vs numeric values
        text_count = 0
        numeric_count = 0
        for val in row:
            if pd.notna(val):
                val_str = str(val).strip()
                if val_str and not val_str.replace('.', '').replace('-', '').isdigit():
                    text_count += 1
                else:
                    numeric_count += 1
        
        text_ratio = text_count / non_null_count if non_null_count > 0 else 0
        
        # Detect header characteristics
        if text_ratio > 0.6:  # Mostly text = likely header
            analysis['header_rows'].append(i)
            
            # Check for merged cells (sparse distribution)
            if text_count > 0 and text_count < non_null_count * 0.4:
                analysis['has_merged_cells'] = True
                
        elif text_ratio < 0.3 and numeric_count > 0:  # Mostly numeric = likely data
            if analysis['data_start_row'] == 0:
                analysis['data_start_row'] = i
            break
    
    # Determine complexity
    analysis['has_multi_level_headers'] = len(analysis['header_rows']) > 1
    analysis['complexity_score'] = (
        len(analysis['header_rows']) * 2 +
        (1 if analysis['has_merged_cells'] else 0) * 3 +
        (1 if analysis['has_multi_level_headers'] else 0) * 2
    )
    
    # Choose processing method
    if analysis['complexity_score'] > 5:
        analysis['recommended_method'] = 'complex'
    elif analysis['has_multi_level_headers']:
        analysis['recommended_method'] = 'multi_level'
    else:
        analysis['recommended_method'] = 'simple'
    
    # Ensure we have at least one header row
    if not analysis['header_rows']:
        analysis['header_rows'] = [0]
    
    return analysis

def _process_complex_excel_structure(file_path: str, sheet_name: str, analysis: dict) -> pd.DataFrame:
    """Process complex Excel structures with advanced techniques"""
    try:
        # Try different header combinations
        for header_combo in [
            analysis['header_rows'],
            [0, 1] if len(analysis['header_rows']) > 1 else [0],
            [0]
        ]:
            try:
                df = pd.read_excel(file_path, sheet_name=sheet_name, header=header_combo)
                
                # Handle MultiIndex columns
                if isinstance(df.columns, pd.MultiIndex):
                    df.columns = _flatten_multiindex_columns(df.columns)
                else:
                    df.columns = _normalize_headers(df.columns)
                
                # Validate the result
                if len(df.columns) > 0 and not df.empty:
                    return df
                    
            except Exception as e:
                logger.debug(f"Header combo {header_combo} failed: {e}")
                continue
        
        # Final fallback
        df = pd.read_excel(file_path, sheet_name=sheet_name, header=0)
        df.columns = _normalize_headers(df.columns)
        return df
        
    except Exception as e:
        logger.error(f"Complex structure processing failed: {e}")
        raise

def _flatten_multiindex_columns(columns: pd.MultiIndex) -> List[str]:
    """Flatten multi-level column names into meaningful single names"""
    flattened = []
    
    for col in columns:
        if isinstance(col, tuple):
            # Join non-null parts with underscore
            parts = []
            for part in col:
                if pd.notna(part) and str(part).strip():
                    parts.append(str(part).strip())
            
            if parts:
                col_name = '_'.join(parts)
            else:
                col_name = f"column_{len(flattened)}"
        else:
            col_name = str(col).strip() if pd.notna(col) else f"column_{len(flattened)}"
        
        flattened.append(_normalize_header(col_name))
    
    return flattened

def _normalize_headers(headers: List[str]) -> List[str]:
    """Normalize a list of headers"""
    return [_normalize_header(str(header)) for header in headers]

def _clean_dataframe(df: pd.DataFrame) -> pd.DataFrame:
    """Clean and validate the processed dataframe"""
    # Remove completely empty rows and columns
    df = df.dropna(how='all').dropna(axis=1, how='all')
    
    # Remove duplicate columns
    df = df.loc[:, ~df.columns.duplicated()]
    
    # Ensure column names are unique
    df.columns = _make_unique_columns(df.columns)
    
    return df

def _make_unique_columns(columns: List[str]) -> List[str]:
    """Ensure all column names are unique"""
    unique_columns = []
    seen = set()
    
    for col in columns:
        original_col = col
        counter = 1
        while col in seen:
            col = f"{original_col}_{counter}"
            counter += 1
        seen.add(col)
        unique_columns.append(col)
    
    return unique_columns

def _normalize_header(header: str) -> str:
    """
    Normalizes a header string to snake_case.
    """
    header = header.strip().lower()
    header = re.sub(r'[^a-z0-9\s_]', '', header) # Remove special characters except underscore
    header = re.sub(r'\s+', '_', header) # Replace spaces with underscores
    return header

def _detect_column_type(series: pd.Series):
    """
    Detects the type of a pandas Series.
    """
    # Try numeric
    try:
        pd.to_numeric(series, errors='raise')
        return "number"
    except (ValueError, TypeError):
        pass
    
    # Try datetime
    try:
        common_date_formats = [
            '%Y-%m-%d %H:%M:%S',
            '%Y-%m-%d %H:%M',
            '%Y-%m-%d',
            '%d-%m-%Y',
            '%m/%d/%Y',
            '%d/%m/%Y',
            '%Y%m%d_%H%M%S', # For formats like '20230707_000000'
            '%Y%m%d',
        ]

        best_converted_series = None
        max_valid_dates = -1

        for fmt in common_date_formats:
            converted_series = pd.to_datetime(series, format=fmt, errors='coerce')
            valid_dates_count = converted_series.count()
            if valid_dates_count > max_valid_dates:
                max_valid_dates = valid_dates_count
                best_converted_series = converted_series
        
        # If a specific format worked for more than 50% of dates, or general coerce works
        if max_valid_dates / len(series) > 0.5:
            return "datetime"
        
        # Fallback to general coercion if no specific format yielded good results but general one might
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)  # Suppress date parsing warnings
            converted_series_general = pd.to_datetime(series, errors='coerce')
        if converted_series_general.count() / len(series) > 0.5:
            return "datetime"
        
    except Exception as e:
        logger.debug(f"Error during date detection for column: {e}")
        pass

    # Default to string
    return "string"

def _process_dataframe(df: pd.DataFrame, file_id: str, file_name: str, sheet_name: str = None) -> List[Document]:
    """
    Processes a pandas DataFrame into a list of Documents (row-level and column summaries).
    """
    documents = []
    normalized_columns = {}
    column_types = {}

    # Normalize headers and detect types
    for col in df.columns:
        normalized_col = _normalize_header(str(col))
        normalized_columns[col] = normalized_col
        column_types[normalized_col] = _detect_column_type(df[col])
    
    df.rename(columns=normalized_columns, inplace=True)

    # Create row-level documents
    for idx, row in df.iterrows():
        row_id = str(uuid.uuid4())
        numeric_fields = {}

        # Dynamically build a descriptive summary for the row
        description_col = str(row.get('description', '')) # Assuming 'description' might exist
        swb_no_col = str(row.get('swb_no', ''))

        row_summary_parts = []
        if description_col: # Start with description if available
            row_summary_parts.append(f"Feeder {description_col.strip()}")
        if swb_no_col:
            row_summary_parts.append(f"location SWB no. {swb_no_col.strip()}")

        # Populate numeric_fields for agent tools
        for col_name, col_type in column_types.items():
            if col_type == "number":
                val = row.get(col_name)
                # Check if val is a Series or scalar
                if isinstance(val, pd.Series):
                    val = val.iloc[0] if not val.empty else np.nan
                
                if pd.notna(val):
                    try:
                        numeric_fields[col_name] = float(val)
                    except (ValueError, TypeError):
                        numeric_fields[col_name] = np.nan # Store NaN for non-convertible numbers
                else:
                    numeric_fields[col_name] = np.nan

        date_columns = [col for col in df.columns if re.match(r'^\\d{{8}}_\\d{{6}}$', col) or re.match(r'^\\d{{2}}-\\d{{2}}-\\d{{4}}$', col)]
        date_columns.sort() # Ensure consistent order

        for i in range(len(date_columns)):
            current_date_col = date_columns[i]
            value_col = current_date_col # The date column itself holds the value
            diff_col = f"difference{i}" if f"difference{i}" in df.columns else None # Find corresponding difference column

            date_val = row.get(current_date_col)
            diff_val = row.get(diff_col)

            if pd.isna(date_val) and pd.isna(diff_val):
                continue

            formatted_date = current_date_col.replace('_', ' ').replace('000000', '').strip()
            
            if not pd.isna(date_val):
                row_summary_parts.append(f"on Date {formatted_date} {str(date_val).strip()} KWH")
            
            if diff_col and not pd.isna(diff_val):
                row_summary_parts.append(f"both day difference {str(diff_val).strip()} KWH")

        # Use SLM to generate a natural language summary for the row (only for important rows)
        # Skip SLM for rows with mostly NaN values or header rows
        non_nan_count = sum(1 for v in row.values if pd.notna(v))
        if non_nan_count >= 3:  # Only process rows with at least 3 non-NaN values
            row_summary_content = _get_slm_summary(row.to_dict(), column_types)
        else:
            # Use simple template-based summary for less important rows
            row_summary_content = " ".join(row_summary_parts) if row_summary_parts else str(row.to_dict())

        metadata = {
            "file_id": file_id,
            "file_name": file_name,
            "sheet": sheet_name,
            "row_id": row_id,
            "columns": list(df.columns),
            "numeric_fields": numeric_fields, # Still collect numeric fields for potential filtering
            "created_at": datetime.now().isoformat(),
            "trusted": TRUSTED_FLAG_DEFAULT,
            "doc_type": "row",
        }
        documents.append(Document(doc_id=str(uuid.uuid4()), file_id=file_id, file_name=file_name, doc_type="row", content=row_summary_content, metadata=metadata))

    # Create column-summary documents
    for col_name in df.columns:
        col_series = df[col_name]
        col_type = column_types[col_name]
        summary_content = f"Column summary for \'{col_name}\'. Data type: {col_type}. "
        
        if col_type == "number":
            if not col_series.empty and col_series.dropna().empty: # If all values are NaN
                min_val_scalar = np.nan
                max_val_scalar = np.nan
                mean_val_scalar = np.nan
            else:
                try:
                    min_val_scalar = float(col_series.min())
                except TypeError:
                    min_val_scalar = np.nan
                try:
                    max_val_scalar = float(col_series.max())
                except TypeError:
                    max_val_scalar = np.nan
                try:
                    mean_val_scalar = float(col_series.mean())
                except TypeError:
                    mean_val_scalar = np.nan

            min_str = f'{min_val_scalar:.2f}' if pd.notna(min_val_scalar) else "N/A"
            max_str = f'{max_val_scalar:.2f}' if pd.notna(max_val_scalar) else "N/A"
            mean_str = f'{mean_val_scalar:.2f}' if pd.notna(mean_val_scalar) else "N/A"

            summary_content += f"Minimum value: {min_str}, Maximum value: {max_str}, Average value: {mean_str}. "
        elif col_type == "datetime":
            # For datetime, provide range if possible
            try:
                min_date = pd.to_datetime(col_series, errors='coerce').min()
                max_date = pd.to_datetime(col_series, errors='coerce').max()
                if pd.notna(min_date) and pd.notna(max_date):
                    summary_content += f"Date range from {min_date.strftime('%Y-%m-%d')} to {max_date.strftime('%Y-%m-%d')}. "
            except Exception:
                pass
        
        unique_sample = col_series.nunique()
        sample_values = col_series.dropna().head(5).values.tolist()
        summary_content += f"Number of unique values: {unique_sample}. Sample values: {sample_values}."
        
        metadata = {
            "file_id": file_id,
            "file_name": file_name,
            "sheet": sheet_name,
            "column_name": col_name,
            "column_type": col_type,
            "created_at": datetime.now().isoformat(),
            "trusted": TRUSTED_FLAG_DEFAULT,
            "doc_type": "column_summary",
        }
        documents.append(Document(doc_id=str(uuid.uuid4()), file_id=file_id, file_name=file_name, doc_type="column_summary", content=summary_content, metadata=metadata))

    return documents

def _chunk_text(text: str, file_id: str, file_name: str, page_no: int = None, section_title: str = None) -> List[Document]:
    """
    Chunks text into smaller documents with overlap.
    """
    # Basic word-based chunking for now
    words = text.split()
    chunks = []
    for i in range(0, len(words), TEXT_CHUNK_SIZE - TEXT_CHUNK_OVERLAP):
        chunk_words = words[i:i + TEXT_CHUNK_SIZE]
        content = " ".join(chunk_words)
        metadata = {
            "file_id": file_id,
            "file_name": file_name,
            "page_no": page_no,
            "section_title": section_title,
            "chunk_id": str(uuid.uuid4()),
            "created_at": datetime.now().isoformat(),
            "trusted": TRUSTED_FLAG_DEFAULT,
            "doc_type": "text_chunk",
        }
        chunks.append(Document(doc_id=str(uuid.uuid4()), file_id=file_id, file_name=file_name, doc_type="text_chunk", content=content, metadata=metadata))
    return chunks

def _process_pdf_file(file_path: str, file_id: str, file_name: str) -> List[Document]:
    """Process PDF file and extract text content"""
    if not PDF_AVAILABLE:
        logger.warning("PDF processing not available. Install PyPDF2 to process PDF files.")
        return []
    
    documents = []
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text.strip():
                    # Create metadata for each page
                    metadata = {
                        "page_number": page_num + 1,
                        "total_pages": len(pdf_reader.pages),
                        "file_type": "pdf",
                        "created_at": datetime.now().isoformat(),
                        "trusted": TRUSTED_FLAG_DEFAULT,
                        "doc_type": "text_chunk",
                    }
                    
                    # Chunk the page text if it's too long
                    if len(text) > TEXT_CHUNK_SIZE:
                        chunks = _chunk_text(text, file_id, file_name)
                        documents.extend(chunks)
                    else:
                        doc = Document(
                            doc_id=str(uuid.uuid4()),
                            file_id=file_id,
                            file_name=file_name,
                            doc_type="text_chunk",
                            content=text,
                            metadata=metadata
                        )
                        documents.append(doc)
        
        logger.info(f"Processed PDF file: {file_name} - {len(documents)} documents created")
        return documents
        
    except Exception as e:
        logger.error(f"Error processing PDF file {file_name}: {e}")
        return []

def _process_docx_file(file_path: str, file_id: str, file_name: str) -> List[Document]:
    """Process DOCX file and extract text content"""
    if not DOCX_AVAILABLE:
        logger.warning("DOCX processing not available. Install python-docx to process DOCX files.")
        return []
    
    documents = []
    try:
        doc = DocxDocument(file_path)
        
        # Extract text from paragraphs
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        combined_text = "\n".join(full_text)
        
        if combined_text.strip():
            # Create metadata
            metadata = {
                "file_type": "docx",
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "created_at": datetime.now().isoformat(),
                "trusted": TRUSTED_FLAG_DEFAULT,
                "doc_type": "text_chunk",
            }
            
            # Chunk the text if it's too long
            if len(combined_text) > TEXT_CHUNK_SIZE:
                chunks = _chunk_text(combined_text, file_id, file_name)
                documents.extend(chunks)
            else:
                doc_obj = Document(
                    doc_id=str(uuid.uuid4()),
                    file_id=file_id,
                    file_name=file_name,
                    doc_type="text_chunk",
                    content=combined_text,
                    metadata=metadata
                )
                documents.append(doc_obj)
        
        logger.info(f"Processed DOCX file: {file_name} - {len(documents)} documents created")
        return documents
        
    except Exception as e:
        logger.error(f"Error processing DOCX file {file_name}: {e}")
        return []

def _process_other_text_file(file_path: str, file_id: str, file_name: str) -> List[Document]:
    """Process other text-based files (log files, etc.)"""
    documents = []
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        text_content = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text_content = f.read()
                break
            except UnicodeDecodeError:
                continue
        
        if text_content is None:
            logger.error(f"Could not decode file {file_name} with any supported encoding")
            return []
        
        # Create metadata
        metadata = {
            "file_type": file_path.split('.')[-1].lower(),
            "encoding": encoding,
            "created_at": datetime.now().isoformat(),
            "trusted": TRUSTED_FLAG_DEFAULT,
            "doc_type": "text_chunk",
        }
        
        # Chunk the text
        chunks = _chunk_text(text_content, file_id, file_name)
        documents.extend(chunks)
        
        logger.info(f"Processed text file: {file_name} - {len(documents)} documents created")
        return documents
        
    except Exception as e:
        logger.error(f"Error processing text file {file_name}: {e}")
        return []

def _process_jsw_energy_report(file_path: str, sheet_name: str, file_context: str = "") -> pd.DataFrame:
    """
    Specialized processor for JSW Energy Reports.
    Handles complex headers, date columns, and merges user context.
    """
    try:
        # Read the sheet with header at row 4 (index 3) based on inspection
        # The inspection showed headers like "Feeder", "SWB No", dates at row 4
        df = pd.read_excel(file_path, sheet_name=sheet_name, header=3)
        
        # Clean column names
        df.columns = [str(col).strip() for col in df.columns]
        
        # Identify key columns
        feeder_col = next((col for col in df.columns if 'Feeder' in col or 'Panel' in col), None)
        swb_col = next((col for col in df.columns if 'SWB' in col), None)
        
        if not feeder_col:
            # Fallback: try reading with header at row 5 (index 4)
            df = pd.read_excel(file_path, sheet_name=sheet_name, header=4)
            df.columns = [str(col).strip() for col in df.columns]
            feeder_col = next((col for col in df.columns if 'Feeder' in col or 'Panel' in col), None)
        
        if not feeder_col:
            logger.warning(f"Could not identify Feeder column in {sheet_name}")
            return _smart_excel_processing(file_path, sheet_name)

        # Filter out rows where Feeder is NaN or "Total"
        df = df[df[feeder_col].notna()]
        df = df[~df[feeder_col].astype(str).str.contains('Total', case=False, na=False)]
        
        # Normalize columns
        normalized_columns = {}
        for col in df.columns:
            if col == feeder_col:
                normalized_columns[col] = 'feeder'
            elif swb_col and col == swb_col:
                normalized_columns[col] = 'swb_no'
            elif 'Unnamed' in col:
                continue # Skip unnamed columns unless they are data
            else:
                # Try to parse date columns
                try:
                    # Check if column is a date-like string or number
                    # The inspection showed dates like "30-06-2024" or floats
                    normalized_columns[col] = str(col)
                except:
                    normalized_columns[col] = _normalize_header(col)
        
        df.rename(columns=normalized_columns, inplace=True)
        
        # Add context column if provided
        if file_context:
            df['file_context'] = file_context
            
        return df

    except Exception as e:
        logger.error(f"JSW Report processing failed for {sheet_name}: {e}")
        return _smart_excel_processing(file_path, sheet_name)

def ingest_file(file_path: str, file_id: str = None, row_limit: Optional[int] = None, file_context: str = "") -> List[Document]:
    """
    Ingests a single file, processes it based on type, and returns a list of Documents.
    """
    file_name = os.path.basename(file_path)
    if file_id is None:
        file_id = str(uuid.uuid4())

    logger.info(f"Starting ingestion for file: {file_path}")
    all_documents: List[Document] = []

    try:
        if file_path.endswith(('.csv', '.xlsx')):
            df = None
            if file_path.endswith('.csv'):
                df = pd.read_csv(file_path)
            elif file_path.endswith('.xlsx'):
                xl = pd.ExcelFile(file_path)
                # Process ALL sheets
                for sheet_name in xl.sheet_names:
                    # Check if it looks like a JSW report (heuristic)
                    is_jsw_report = "Energy" in file_name or "Report" in file_name
                    
                    if is_jsw_report:
                        df_sheet = _process_jsw_energy_report(file_path, sheet_name, file_context)
                    else:
                        df_sheet = _smart_excel_processing(file_path, sheet_name, row_limit)
                    
                    if df_sheet is not None and not df_sheet.empty:
                        # Add sheet name to context if not already present
                        sheet_context = f"Sheet: {sheet_name}. {file_context}"
                        all_documents.extend(_process_dataframe(df_sheet, file_id, file_name, sheet_name=sheet_name))
            
            if df is not None: # For CSVs that are processed directly
                if row_limit: # Apply row limit if provided
                    df = df.head(row_limit)
                all_documents.extend(_process_dataframe(df, file_id, file_name))

        elif file_path.endswith('.pdf'):
            all_documents.extend(_process_pdf_file(file_path, file_id, file_name))
            log_process_completion(f"Ingestion of PDF: {file_name}", details="Processed as PDF document")

        elif file_path.endswith('.docx'):
            all_documents.extend(_process_docx_file(file_path, file_id, file_name))
            log_process_completion(f"Ingestion of DOCX: {file_name}", details="Processed as DOCX document")

        elif file_path.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            # Prepend context to text content
            if file_context:
                text_content = f"Context: {file_context}\n\n{text_content}"
            all_documents.extend(_chunk_text(text_content, file_id, file_name))
            log_process_completion(f"Ingestion of TXT: {file_name}", details="Processed as textual data")

        elif file_path.endswith(('.log', '.md', '.json', '.xml', '.yaml', '.yml')):
            all_documents.extend(_process_other_text_file(file_path, file_id, file_name))
            log_process_completion(f"Ingestion of {file_name.split('.')[-1].upper()}: {file_name}", details="Processed as text document")

        else:
            logger.warning(f"Unsupported file type for ingestion: {file_name}")
            log_process_completion(f"Ingestion of {file_name}", status="skipped", details="Unsupported file type")
            return []

        logger.info(f"Successfully ingested {len(all_documents)} documents from {file_name}")
        return all_documents

    except Exception as e:
        logger.error(f"Error ingesting file {file_name}: {e}", exc_info=True)
        log_process_completion(f"Ingestion of {file_name}", status="failed", details=str(e))
        return []
